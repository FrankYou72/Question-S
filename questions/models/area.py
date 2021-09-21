from django.db import models
from django.db.models import manager, CASCADE, ForeignKey, JSONField, FileField
from django.db.models.fields import CharField, TextField
from random import choice
from datetime import date, datetime
from docx import Document


class Area(models.Model):
    area = CharField(max_length=100)
    objects = manager.Manager()

    def get_list(self, n):
        temas = self.questionmodel_set.all()
        question_list = {
            'questões': {},
            'gabarito': {}
        }

        for i in range(n):
            q_model = choice(temas)
            question = q_model.get_question()
            question_list['questões'][str(i+1)] = question.enunciado
            question_list['gabarito'][str(i+1)] = question.gabarito

        return QuestionList(area=self, enunciados=question_list['questões'], gabarito=question_list['gabarito'])

    def __str__(self):
        return self.area


class QuestionList(models.Model):
    area = ForeignKey(Area, on_delete=CASCADE)
    enunciados = JSONField()
    gabarito = JSONField()
    #docx_file = FileField(null=True)

    def to_docx(self):
        hoje = date.today()
        agora = datetime.now()
        Lista = Document()
        cabecalho = f'Lista de Exercícios de {self.area.area}'
        Lista.add_heading(cabecalho, level=0)
        p = 1
        path = "C:/Users/Franklin/OneDrive/repos/QuestionS_Web_Project/questions/static/text/"

        for e in self.enunciados.keys():
            Lista.add_heading(f'Questão {e}', level=2)
            Lista.add_paragraph(self.enunciados[e])
            p += 1

        Lista.add_page_break()
        Lista.add_heading('GABARITO')

        for g in self.gabarito.keys():
            Lista.add_paragraph(f'{g} ---> {self.gabarito[g]}')

        Lista.add_page_break()
        #self.docx_file = Lista.save(path + f'{hoje} - {agora.hour}-{agora.min}-{agora.second}.docx')

