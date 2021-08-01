import random
from sympy import symbols, Eq, solve, sin, cos, sqrt, pi, tan, init_printing, acos, asin, atan, log, exp
from time import sleep
from .Q_data import Packs, opes
from docx import Document
from datetime import datetime, date
from math import ceil

global hoje
global agora
global sig_fig


hoje = date.today()
agora = datetime.now()
sig_fig = 1

class Question:
    def __init__(self, area, tema, sig_fig = 2):
        self.area = area
        self.tema = tema
        self.equacao = Packs[self.area][1][self.tema][0]
        self.enunciado = Packs[self.area][1][self.tema][1]
        self.variaveis = get_vars(self.equacao, Packs[self.area][0])
        self.grandezas = {}
        for g in Packs[self.area][2].keys():
            if str(g) in self.variaveis:
                self.grandezas[g] = Packs[self.area][2][g]
        self.unidades = {}
        for u in Packs[self.area][3].keys():
            if str(u) in self.variaveis:
                self.unidades[u] = Packs[self.area][3][u]
        self.equacoes = []
        for v in self.variaveis: #Return the equation for each variable
            s = solve(self.equacao, v)
            if str(s) == '[]':
               pass
            else:
                self.equacoes.append(s)
        self.sig_fig = sig_fig

    def generate(self):
        global opes
        dados = self.variaveis
        escolha = str(random.choice(self.equacoes))
        escolha_copia = escolha
        for o in opes:
            if o in str(escolha_copia):
                escolha = escolha.replace(o, ' ')
        for v in self.variaveis:
            if str(v) not in escolha.split():
                icognita = v
                dados.remove(v)
        escolha = escolha.strip('[]')
        valores = {}
        for d in dados: #Parâmetros para valores de variável
            for u in self.unidades.keys():
                if str(d) == str(u):
                    if len(self.unidades[u][1]) == 3:
                        A = self.unidades[u][1][0]
                        B = self.unidades[u][1][1]
                        C = self.unidades[u][1][2]
                        valores[d] = random.randint(A, B)/C
                    elif len(self.unidades[u][1]) == 5:
                        A = self.unidades[u][1][0]
                        B = self.unidades[u][1][1]
                        C = self.unidades[u][1][2]
                        D = self.unidades[u][1][3]
                        E = self.unidades[u][1][4]
                        valores[d] = (random.randint(A, B)/C)*10**(random.randint(D, E))
        for g in self.grandezas.keys():
            if str(icognita) == str(g):
                ic = self.grandezas[g]
        escolha_copia = escolha
        for o in opes:
            if o in escolha:
                escolha_copia = escolha_copia.replace(o, f' {o} ')
        escolha = escolha_copia.split()
        valores_keys = []
        for k in valores.keys():
            valores_keys.append(k)
        for v in valores_keys:
            for e in escolha:
                if str(v) == str(e):
                    i = escolha.index(e)
                    escolha.remove(e)
                    escolha.insert(i, str(valores[v]))
        escolha = ''.join(escolha)

        E = self.enunciado + '\n'
        for k in valores.keys():
            for g in self.grandezas.keys():
                if g == k:
                    E += self.grandezas[g] + ' = ' + str(valores[k]) + ' ' + str(self.unidades[g][0]) + ',\n '
        E += 'assim, calcule o valor do(a)\n' + ic + '\n'

        try:
            resposta = eval(escolha)
            if type(resposta) == tuple or type(resposta) == list:
                print('resposta composta:', resposta)
                for r in resposta:
                    if str(r)[0] != '-':
                        resposta = float(r)
                        break
            Gabarito = round(float(resposta), self.sig_fig)
        except:
            Gabarito = 5*'-=' + 'Sem solução Real' +  5*'=-'
        
        return {'Enunciado':E, 'Gabarito':Gabarito}

    def __str__(self):
        return self.area + '/' + self.tema


def get_vars(eq, variables):
    global opes
    argus = str(eq.args[0]) #Gets the math expression in string
    for o in opes: #Will loop through the expression and remove all operators, leaving only the variables
        if o in argus:
            argus = argus.replace(o, ' ')
    vari = argus.split() #Returns a list with the string variables
    vars = []
    for v in vari: #Get the Sympy Symbols for the selected variables
        for k in variables.keys():
            if v == str(k):
                vars.append(str(k))
    return vars

def get_self_quantities(eq, variables, quantities):
    vs = get_vars(eq, variables)
    qt = {}
    for q in quantities:
        for v in vs:
            if str(q) == v:
                qt[str(v)] = quantities[q]
    return qt

def get_self_units(eq, variables, units):
    vs = get_vars(eq, variables)
    un = {}
    for u in units:
        for v in vs:
            if str(u) == v:
                un[v] = units[u]
    return un

def get_str_eqs(eq, variables):
    global opes
    argus = str(eq.args[0]) #Gets the math expression in string
    for o in opes: #Will loop through the expression and remove all operators, leaving only the variables
        if o in argus:
            argus = argus.replace(o, ' ')
    vari = argus.split() #Returns a list with the string variables
    vars = []
    for v in vari: #Get the Sympy Symbols for the selected variables
        for k in variables.keys():
            if v == str(k) and k not in vars:
                vars.append(k)
    equations = []
    for v in vars: #Return the equation for each variable
        s = solve(eq, v)
        if str(s) == '[]':
            pass
        else:
            equations.append(str(s))
    return equations

def get_equations(eq, variables):
    global opes
    argus = str(eq.args[0]) #Gets the math expression in string
    for o in opes: #Will loop through the expression and remove all operators, leaving only the variables
        if o in argus:
            argus = argus.replace(o, ' ')
    vari = argus.split() #Returns a list with the string variables
    vars = []
    for v in vari: #Get the Sympy Symbols for the selected variables
        for k in variables.keys():
            if v == str(k) and k not in vars:
                vars.append(k)
    equations = []
    print('Variables as received from variables function: ', vars)
    for v in vars: #Return the equation for each variable
        s = solve(eq, v)
        if str(s) == '[]':
            pass
        else:
            equations.append(s)
    return equations

def get_figures(eq ,eqs, quantities, variables, units):
    global opes
    vars = get_vars(eq, variables)
    Pick = random.choice(eqs)
    print('Equations as received from Equations function: ', eqs)
    PV = str(Pick)
    for o in opes:
        if o in str(Pick):
            PV = PV.replace(o, ' ')
    for vv in vars:
        if str(vv) not in PV.split():
            solve_for = vv
            vars.remove(vv)
    values = {}
    for vv in vars: #Parâmetros para valores de variável
        for un in units.keys():
            if str(vv) == str(un):
                if len(units[un][1]) == 3:
                    A = units[un][1][0]
                    B = units[un][1][1]
                    C = units[un][1][2]
                    values[vv] = random.randint(A, B)/C
                elif len(units[un][1]) == 5:
                    A = units[un][1][0]
                    B = units[un][1][1]
                    C = units[un][1][2]
                    D = units[un][1][3]
                    E = units[un][1][4]
                    values[vv] = (random.randint(A, B)/C)*10**(random.randint(D, E))
                pass
    return Pick, solve_for, values

def make_question(eq, quantities, units, variables, prompt, sig_fig):
    global opes
    equations = get_equations(eq, variables)
    figs = get_figures(eq, equations, quantities, variables, units)
    Pick = figs[0]
    print('figures returned from Figures function: ', figs)
    print('Pick as received for question function: ', Pick)
    pick = str(Pick).strip('[]')
    solve_for = figs[1]
    for q in quantities.keys():
        if str(solve_for) == str(q):
            sf = quantities[q]
    print('sf : ', sf)
    New_Pick = pick
    for o in opes:
        if o in New_Pick:
            New_Pick = New_Pick.replace(o, f' {o} ')
    vals = figs[2]
    pick = New_Pick.split()
    vals_keys = []
    for k in vals.keys():
        vals_keys.append(k)
    for vv in vals_keys:
        for pp in pick:
            if str(vv) == str(pp):
                i = pick.index(pp)
                pick.remove(pp)
                pick.insert(i, str(vals[vv]))
    pick = ''.join(pick)

    Q = prompt + '\n'
    for k in vals.keys():
        for q in quantities.keys():
            if q == k:
                Q += quantities[q] + ' = ' + str(vals[k]) + ' ' + str(units[q][0]) + ',\n '
    Q += 'assim, calcule o valor do(a)\n' + sf + '\n'

    if '*I' in pick:
        pick = pick.replace('*I' , 'j')    
    if 'I*' in pick:
        pick = pick.replace('I*', 'j')

    try:
        resposta = eval(pick)
        if type(resposta) == tuple or type(resposta) == list:
            print('resposta composta:', resposta)
            for r in resposta:
                if str(r)[0] != '-':
                    resposta = float(r)
                    print('resposta escolhida:', r )
                    break
        Gabarito = round(float(resposta), sig_fig)
    except:
        print('pick at except: ', pick)
        print('resposta: ', eval(pick))
        Gabarito = 5*'-=' + 'Sem solução Real' +  5*'=-'

    return Q, f"Gabarito: {Gabarito} {units[solve_for][0]}"

def list_ex(n, variables, equations, quantities, units, subject, sigs=2):
    hoje = date.today()
    agora = datetime.now()
    Lista = Document()
    cabecalho = f'Lista de Exercícios de {subject}'
    Lista.add_heading(cabecalho, level=0)
    names = list(equations.keys())
    gabarito = []
    p = 1

    for i in range(n):
        eq_family = equations[random.choice(names)]
        equation = eq_family[0]
        prompt = eq_family[1]
        perg = make_question(equation, quantities, units, variables, prompt, sigs)
        gabarito.append(perg[1])
        Lista.add_heading(f'Questão {p}', level=2)
        Lista.add_paragraph(perg[0])
        p += 1

    Lista.add_page_break()
    p = 1
    Lista.add_heading('GABARITO')

    for g in gabarito:
        Lista.add_paragraph(f'{p} ---> {g}')
        p += 1
    Lista.add_page_break()
    Lista.save(f'ListEx {hoje.year}-{hoje.month}-{hoje.day}-{agora.hour}-{agora.minute}-{agora.second}.docx')

def generate_custom(question_list, gabarito):
    Lista = Document()
    cabecalho = f'Lista de Exercícios Personalizada'
    Lista.add_heading(cabecalho, level=0)

    for i in range(len(question_list)):
        perg = question_list[i]
        Lista.add_heading(f'Questão {i + 1}', level=2)
        Lista.add_paragraph(perg)

    Lista.add_page_break()
    p = 1
    Lista.add_heading('GABARITO')

    for i in range(len(gabarito)):
        Lista.add_paragraph(f'{i + 1} ---> {gabarito[i]}')
    Lista.add_page_break()
    Lista.save(f'CustEx {hoje.year}-{hoje.month}-{hoje.day}-{agora.hour}-{agora.minute}-{agora.second}.docx')

def generate_sing(vars, eq, qua, uns, sigs = 2):
    gabarito = ''
    p = 1
    equation = eq[0]
    prompt = eq[1]
    perg = make_question(equation, qua, uns, vars, prompt, sigs)
    gabarito = str(perg[1])
    quest = perg[0]
    return quest, gabarito